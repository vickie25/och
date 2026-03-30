"""
Certificate DOCX Generation Service
Generates Word (.docx) certificates by filling templates with JSON data.
"""
import os
import io
import re
from datetime import datetime
from typing import Dict, Any, Optional, List

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False

from django.conf import settings
from django.core.files.base import ContentFile


class CertificateDOCXGenerator:
    """Generate .docx certificates by filling templates with JSON data."""
    
    # Default placeholder mappings (JSON field -> template placeholder) per OCH spec v1.0
    DEFAULT_PLACEHOLDERS = {
        # Student info (6.2 spec)
        'student_name': '{{student_name}}',
        'student_id': '{{student_id}}',
        
        # Certificate details (6.2 spec)
        'completion_date': '{{completion_date}}',
        'expiry_date': '{{expiry_date}}',
        'certificate_id': '{{certificate_id}}',
        'certificate_status': '{{certificate_status}}',
        'renewal_count': '{{renewal_count}}',
        'total_hours': '{{total_hours}}',
        'missions_completed': '{{missions_completed}}',
        
        # Issuer info (6.2 spec)
        'issuer_name': '{{issuer_name}}',
        'issuer_title': '{{issuer_title}}',
        'signature_image': '{{signature_image}}',
        
        # Verification (6.2 spec)
        'verification_qr': '{{verification_qr}}',
        'accreditation_badge': '{{accreditation_badge}}',
        
        # Cohort info (6.2 spec)
        'cohort_name': '{{cohort_name}}',
        'cohort_description': '{{cohort_description}}',
        
        # Legacy placeholders (for backward compatibility)
        'first_name': '{{FIRST_NAME}}',
        'last_name': '{{LAST_NAME}}',
        'student_email': '{{STUDENT_EMAIL}}',
        'program_name': '{{PROGRAM_NAME}}',
        'course_name': '{{COURSE_NAME}}',
        'track_name': '{{TRACK_NAME}}',
        'program_duration': '{{PROGRAM_DURATION}}',
        'completion_hours': '{{COMPLETION_HOURS}}',
        'issue_date': '{{ISSUE_DATE}}',
        'expiration_date': '{{EXPIRATION_DATE}}',
        'director_name': '{{DIRECTOR_NAME}}',
        'institution_name': '{{INSTITUTION_NAME}}',
        'skills': '{{SKILLS}}',
        'competencies': '{{COMPETENCIES}}',
        'grade': '{{GRADE}}',
        'score': '{{SCORE}}',
    }
    
    # Template file paths
    TEMPLATE_DIR = os.path.join(settings.MEDIA_ROOT, 'certificates', 'templates')
    OUTPUT_DIR = os.path.join(settings.MEDIA_ROOT, 'certificates', 'generated')
    
    # Track/Level to template mapping
    # Maps track keys to their template prefixes
    TRACK_TEMPLATE_MAP = {
        'DEF': 'OCH_Certificate_DEF',      # Defensive Track
        'INN': 'OCH_Certificate_INN',      # Innovation
        'GRC': 'OCH_Certificate_GRC',      # Governance Risk and Compliance
        'OFF': 'OCH_Certificate_OFF',      # Offensive Track
        'L0': 'OCH_Certificate_L0',        # Foundation Track
    }
    
    # Level mapping (standardizes level names)
    LEVEL_MAP = {
        'beginner': 'Beginner',
        'intermediate': 'Intermediate',
        'advanced': 'Advanced',
        'mastery': 'Mastery',
        'L1': 'Beginner',
        'L2': 'Intermediate',
        'L3': 'Advanced',
        'L4': 'Mastery',
    }
    
    # Cohort template for cohort-based certificates
    COHORT_TEMPLATE = 'OCH_Certificate_Cohort'
    
    @classmethod
    def fill_template(cls, template_path: str, data: Dict[str, Any]):
        """
        Fill a .docx template with JSON data.
        
        Args:
            template_path: Path to the .docx template file
            data: Dictionary containing certificate data
            
        Returns:
            Filled Document object
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
        
        # Load template
        doc = Document(template_path)
        
        # Process all paragraphs
        for paragraph in doc.paragraphs:
            cls._replace_placeholders_in_paragraph(paragraph, data)
        
        # Process all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        cls._replace_placeholders_in_paragraph(paragraph, data)
        
        # Process headers/footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for paragraph in header.paragraphs:
                        cls._replace_placeholders_in_paragraph(paragraph, data)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for paragraph in footer.paragraphs:
                        cls._replace_placeholders_in_paragraph(paragraph, data)
        
        return doc
    
    @classmethod
    def _replace_placeholders_in_paragraph(cls, paragraph, data: Dict[str, Any]):
        """Replace placeholders in a single paragraph."""
        for run in paragraph.runs:
            text = run.text
            for key, value in data.items():
                placeholder = cls.DEFAULT_PLACEHOLDERS.get(key, f'{{{{{key.upper()}}}}}')
                if placeholder in text:
                    # Format value based on type
                    formatted_value = cls._format_value(value)
                    text = text.replace(placeholder, formatted_value)
            run.text = text
        
        # Also check the paragraph text as a whole (for multi-run placeholders)
        full_text = paragraph.text
        original_text = full_text
        for key, value in data.items():
            placeholder = cls.DEFAULT_PLACEHOLDERS.get(key, f'{{{{{key.upper()}}}}}')
            if placeholder in full_text:
                formatted_value = cls._format_value(value)
                full_text = full_text.replace(placeholder, formatted_value)
        
        # If we made changes, update the paragraph
        if full_text != original_text:
            paragraph.clear()
            run = paragraph.add_run(full_text)
            run.font.size = Pt(12)
    
    @classmethod
    def _format_value(cls, value: Any) -> str:
        """Format a value for insertion into the document."""
        if value is None:
            return ''
        if isinstance(value, datetime):
            return value.strftime('%B %d, %Y')
        if isinstance(value, list):
            return ', '.join(str(item) for item in value)
        return str(value)
    
    @classmethod
    def get_template_for_track_level(cls, track_key: str, level: str = None, is_cohort: bool = False) -> str:
        """
        Get the correct template name for a track and level combination.
        
        Args:
            track_key: Track code (DEF, INN, GRC, OFF, L0)
            level: Level name (Beginner, Intermediate, Advanced, Mastery) or None
            is_cohort: Whether this is a cohort-based certificate
            
        Returns:
            Template name (without .docx extension)
        """
        if is_cohort:
            return cls.COHORT_TEMPLATE
        
        # Normalize track key
        track_key = track_key.upper() if track_key else 'L0'
        
        # Get template prefix for track
        template_prefix = cls.TRACK_TEMPLATE_MAP.get(track_key)
        
        if not template_prefix:
            # Fallback to default if track not found
            return 'default'
        
        # Normalize level
        if level:
            level_normalized = cls.LEVEL_MAP.get(level.lower(), level.title())
        else:
            level_normalized = 'Beginner'
        
        # Build template name: OCH_Certificate_DEF_Beginner
        template_name = f"{template_prefix}_{level_normalized}"
        
        # Check if template exists
        template_path = os.path.join(cls.TEMPLATE_DIR, f"{template_name}.docx")
        if os.path.exists(template_path):
            return template_name
        
        # Try fallback: OCH_Certificate_DEF (without level)
        template_name = template_prefix
        template_path = os.path.join(cls.TEMPLATE_DIR, f"{template_name}.docx")
        if os.path.exists(template_path):
            return template_name
        
        # Final fallback to default
        return 'default'
    
    @classmethod
    def generate_certificate_docx(cls, certificate, template_name: str = None) -> bytes:
        """
        Generate a .docx certificate from template and certificate data.
        
        Args:
            certificate: Certificate model instance
            template_name: Optional template name. If not provided, auto-detected from track/level.
            
        Returns:
            Bytes of the generated .docx file
        """
        # Auto-detect template if not provided
        if not template_name:
            template_name = cls._resolve_template_for_certificate(certificate)
        
        # Build template path
        template_filename = f"{template_name}.docx"
        template_path = os.path.join(cls.TEMPLATE_DIR, template_filename)
        
        # Check if template exists, otherwise use default
        if not os.path.exists(template_path):
            # Create default template if it doesn't exist
            template_path = cls._create_default_template()
            template_name = 'default'
        
        # Build data dictionary from certificate
        data = cls._build_certificate_data(certificate)
        
        # Fill template
        doc = cls.fill_template(template_path, data)
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
    
    @classmethod
    def _resolve_template_for_certificate(cls, certificate) -> str:
        """Resolve the correct template name for a certificate based on its track and level."""
        enrollment = certificate.enrollment
        cohort = enrollment.cohort
        track = cohort.track
        
        # Get track key
        track_key = None
        if hasattr(track, 'key'):
            track_key = track.key
        elif hasattr(track, 'code'):
            track_key = track.code
        else:
            # Extract from track name
            track_name = track.name.upper()
            if 'DEF' in track_name or 'DEFENSIVE' in track_name:
                track_key = 'DEF'
            elif 'OFF' in track_name or 'OFFENSIVE' in track_name:
                track_key = 'OFF'
            elif 'INN' in track_name or 'INNOVATION' in track_name:
                track_key = 'INN'
            elif 'GRC' in track_name or 'GOVERNANCE' in track_name:
                track_key = 'GRC'
            elif 'L0' in track_name or 'FOUNDATION' in track_name:
                track_key = 'L0'
        
        # Determine level from enrollment or track
        level = None
        if hasattr(enrollment, 'level'):
            level = enrollment.level
        elif hasattr(track, 'level'):
            level = track.level
        elif hasattr(cohort, 'level'):
            level = cohort.level
        
        # Check if cohort-based certificate
        is_cohort = certificate.template_used == 'cohort' if hasattr(certificate, 'template_used') else False
        
        return cls.get_template_for_track_level(track_key, level, is_cohort)
    
    @classmethod
    def _build_certificate_data(cls, certificate) -> Dict[str, Any]:
        """Build data dictionary from certificate model per OCH spec v1.0."""
        enrollment = certificate.enrollment
        user = enrollment.user
        cohort = enrollment.cohort
        track = cohort.track
        program = track.program
        
        # Generate formatted certificate ID: OCH-{TRACK}-{LEVEL}-{YYYYMMDD}-{RANDOM}
        date_str = certificate.issue_date.strftime('%Y%m%d') if certificate.issue_date else datetime.now().strftime('%Y%m%d')
        random_suffix = str(certificate.id)[:8].upper()
        track_key = track.key.upper() if hasattr(track, 'key') else 'GEN'
        cert_id_formatted = f"OCH-{track_key}-L1-{date_str}-{random_suffix}"
        
        return {
            # Student info (6.2 spec)
            'student_name': f"{user.first_name} {user.last_name}".strip() or user.email,
            'student_id': str(user.id),
            
            # Certificate details (6.2 spec)
            'completion_date': enrollment.completed_at or enrollment.joined_at,
            'expiry_date': certificate.expiry_date,
            'certificate_id': cert_id_formatted,
            'certificate_status': certificate.status.upper(),
            'renewal_count': str(certificate.renewal_count),
            'total_hours': str(certificate.total_hours),
            'missions_completed': str(certificate.missions_completed),
            
            # Issuer info (6.2 spec)
            'issuer_name': program.director.get_full_name() if program.director else 'Program Director',
            'issuer_title': 'Program Director',
            'signature_image': '',  # Placeholder - loaded from platform config
            
            # Verification (6.2 spec)
            'verification_qr': '',  # Generated per certificate
            'accreditation_badge': '',  # Platform config
            
            # Cohort info (6.2 spec)
            'cohort_name': cohort.name,
            'cohort_description': cohort.description or '',
            
            # Legacy placeholders (backward compatibility)
            'first_name': user.first_name or '',
            'last_name': user.last_name or '',
            'student_email': user.email,
            'program_name': program.name,
            'course_name': track.name,
            'track_name': track.name,
            'program_duration': f"{cohort.start_date} to {cohort.end_date}" if cohort.start_date and cohort.end_date else 'N/A',
            'completion_hours': str(certificate.total_hours),
            'issue_date': certificate.issue_date,
            'expiration_date': certificate.expiry_date,
            'director_name': program.director.get_full_name() if program.director else 'Program Director',
            'institution_name': 'Ongoza Cyber Hub',
            'skills': 'Cybersecurity, Network Security, Incident Response',
            'competencies': 'Technical Proficiency, Problem Solving, Critical Thinking',
            'grade': enrollment.grade or 'Pass',
            'score': enrollment.grade or 'Pass',
        }
    
    @classmethod
    def _create_default_template(cls) -> str:
        """Create a default certificate template if none exists."""
        os.makedirs(cls.TEMPLATE_DIR, exist_ok=True)
        
        template_path = os.path.join(cls.TEMPLATE_DIR, 'default.docx')
        
        if os.path.exists(template_path):
            return template_path
        
        # Create a basic template
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('CERTIFICATE')
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 64, 175)
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('OF COMPLETION')
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(100, 116, 139)
        
        doc.add_paragraph()
        
        # Body
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p1.add_run('This certifies that')
        run.font.size = Pt(14)
        
        doc.add_paragraph()
        
        # Student name
        name = doc.add_paragraph()
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name.add_run('{{STUDENT_NAME}}')
        run.font.size = Pt(28)
        run.font.bold = True
        
        doc.add_paragraph()
        
        # Has completed
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p2.add_run('has successfully completed the')
        run.font.size = Pt(14)
        
        doc.add_paragraph()
        
        # Program name
        program = doc.add_paragraph()
        program.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = program.add_run('{{PROGRAM_NAME}}')
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 64, 175)
        
        doc.add_paragraph()
        
        # Details
        details = doc.add_paragraph()
        details.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = details.add_run('Track: {{TRACK_NAME}} | Cohort: {{COHORT_NAME}}')
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(100, 116, 139)
        
        doc.add_paragraph()
        
        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run('Completed on {{COMPLETION_DATE}}')
        run.font.size = Pt(14)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Certificate ID
        cert_id = doc.add_paragraph()
        cert_id.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cert_id.add_run('Certificate ID: {{CERTIFICATE_ID}}')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(148, 163, 184)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Signature section
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Director signature
        cell1 = sig_table.rows[0].cells[0]
        cell1.paragraphs[0].add_run('_________________').font.size = Pt(12)
        cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        cell2 = sig_table.rows[1].cells[0]
        cell2.paragraphs[0].add_run('{{DIRECTOR_NAME}}').font.size = Pt(12)
        cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        cell3 = sig_table.rows[0].cells[1]
        cell3.paragraphs[0].add_run('_________________').font.size = Pt(12)
        cell3.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        cell4 = sig_table.rows[1].cells[1]
        cell4.paragraphs[0].add_run('Date').font.size = Pt(12)
        cell4.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save template
        doc.save(template_path)
        
        return template_path
    
    @classmethod
    def docx_to_pdf(cls, docx_bytes: bytes) -> bytes:
        """Convert DOCX bytes to PDF bytes."""
        if not DOCX2PDF_AVAILABLE:
            raise ImportError("docx2pdf is required for conversion. Install with: pip install docx2pdf")
        
        # Save DOCX to temp file
        import tempfile
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
            tmp_docx.write(docx_bytes)
            tmp_docx_path = tmp_docx.name
        
        # Convert to PDF
        tmp_pdf_path = tmp_docx_path.replace('.docx', '.pdf')
        convert(tmp_docx_path, tmp_pdf_path)
        
        # Read PDF bytes
        with open(tmp_pdf_path, 'rb') as f:
            pdf_bytes = f.read()
        
        # Cleanup
        os.unlink(tmp_docx_path)
        os.unlink(tmp_pdf_path)
        
        return pdf_bytes
    
    @classmethod
    def list_available_templates(cls) -> List[str]:
        """List available certificate templates."""
        if not os.path.exists(cls.TEMPLATE_DIR):
            return []
        
        templates = []
        for filename in os.listdir(cls.TEMPLATE_DIR):
            if filename.endswith('.docx'):
                templates.append(filename.replace('.docx', ''))
        
        return templates
    
    @classmethod
    def get_template_placeholders(cls, template_name: str) -> List[str]:
        """Get list of placeholders in a template."""
        template_path = os.path.join(cls.TEMPLATE_DIR, f"{template_name}.docx")
        
        if not os.path.exists(template_path):
            return []
        
        if not DOCX_AVAILABLE:
            return []
        
        doc = Document(template_path)
        placeholders = set()
        
        # Extract placeholders from all text
        placeholder_pattern = r'\{\{([^}]+)\}\}'
        
        for paragraph in doc.paragraphs:
            matches = re.findall(placeholder_pattern, paragraph.text)
            placeholders.update(matches)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = re.findall(placeholder_pattern, paragraph.text)
                        placeholders.update(matches)
        
        return sorted(list(placeholders))


# Convenience function
def generate_certificate_docx(certificate, template_name='default') -> bytes:
    """Generate a .docx certificate for the given certificate instance."""
    return CertificateDOCXGenerator.generate_certificate_docx(certificate, template_name)
